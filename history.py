from __future__ import annotations

import csv
import json
import time
from pathlib import Path
from typing import Any

from storage import DATA_DIR, load_history, save_history


# 历史记录用于复盘、导出和后续纠错优化。


def add_history(record: dict[str, Any]) -> None:
    records = load_history()
    record["id"] = f"{int(time.time() * 1000)}"
    record["timestamp"] = int(time.time())
    records.insert(0, record)
    save_history(records[:2000])


def clear_history() -> None:
    # 清空历史只影响 data/history.json，不会改动配置方案。
    save_history([])


def export_json(path: Path | None = None) -> Path:
    records = load_history()
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.json"
    target.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def export_csv(path: Path | None = None) -> Path:
    records = load_history()
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.csv"
    with target.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["时间", "方案", "模式", "AI分数", "最终分数", "识别答案", "评语", "是否纠错", "小题得分", "勤勉加分", "双评结果"])
        for r in records:
            dual = r.get("dual_eval") or {}
            sub_scores = "; ".join(f"{x.get('label')}={x.get('score')}/{x.get('maxScore')}" for x in r.get("sub_scores", []) or [])
            writer.writerow([
                time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(r.get("timestamp", 0))),
                r.get("preset", ""),
                r.get("mode", ""),
                r.get("ai_score", ""),
                r.get("final_score", ""),
                r.get("student_answer", ""),
                r.get("comment", ""),
                "是" if r.get("corrected") else "否",
                sub_scores,
                r.get("bonus", ""),
                dual.get("result", ""),
            ])
    return target


def export_html(path: Path | None = None) -> Path:
    records = load_history()
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.html"
    rows = []
    for r in records:
        rows.append(
            "<tr>"
            f"<td>{time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(r.get('timestamp', 0)))}</td>"
            f"<td>{_esc(r.get('preset', ''))}</td>"
            f"<td>{_esc(r.get('mode', ''))}</td>"
            f"<td>{_esc(r.get('ai_score', ''))}</td>"
            f"<td>{_esc(r.get('final_score', ''))}</td>"
            f"<td>{_esc(r.get('student_answer', ''))}</td>"
            f"<td>{_esc(r.get('comment', ''))}</td>"
            f"<td>{_esc(_sub_score_text(r.get('sub_scores')))}</td>"
            f"<td>{_esc((r.get('dual_eval') or {}).get('result', ''))}</td>"
            "</tr>"
        )
    html = """<!doctype html><html><head><meta charset="utf-8"><title>评阅历史</title>
<style>body{font-family:Microsoft YaHei UI,Segoe UI,sans-serif;margin:24px;color:#172033}table{border-collapse:collapse;width:100%}td,th{border:1px solid #d8dee8;padding:8px;vertical-align:top}th{background:#f3f6fa}</style>
</head><body><h1>评阅历史</h1><table><thead><tr><th>时间</th><th>方案</th><th>模式</th><th>AI分数</th><th>最终分数</th><th>识别答案</th><th>评语</th><th>小题得分</th><th>双评结果</th></tr></thead><tbody>""" + "\n".join(rows) + "</tbody></table></body></html>"
    target.write_text(html, encoding="utf-8")
    return target


def export_docx(path: Path | None = None) -> Path:
    records = load_history()
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return export_html(target.with_suffix(".html"))
    doc = Document()
    doc.add_heading("评阅历史", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["时间", "方案", "模式", "AI分数", "最终分数", "识别答案", "评语"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for r in records:
        cells = table.add_row().cells
        values = [
            time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(r.get("timestamp", 0))),
            r.get("preset", ""),
            r.get("mode", ""),
            r.get("ai_score", ""),
            r.get("final_score", ""),
            r.get("student_answer", ""),
            r.get("comment", ""),
        ]
        for cell, value in zip(cells, values):
            cell.text = str(value)
    doc.save(target)
    return target


def export_xlsx(path: Path | None = None) -> Path:
    records = load_history()
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.xlsx"
    try:
        from openpyxl import Workbook  # type: ignore
    except ImportError:
        return export_csv(target.with_suffix(".csv"))
    wb = Workbook()
    ws = wb.active
    ws.title = "评阅历史"
    headers = ["时间", "方案", "模式", "AI分数", "最终分数", "识别答案", "评语", "是否纠错", "小题得分", "勤勉加分", "双评结果"]
    ws.append(headers)
    for r in records:
        dual = r.get("dual_eval") or {}
        ws.append([
            time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(r.get("timestamp", 0))),
            r.get("preset", ""),
            r.get("mode", ""),
            r.get("ai_score", ""),
            r.get("final_score", ""),
            r.get("student_answer", ""),
            r.get("comment", ""),
            "是" if r.get("corrected") else "否",
            _sub_score_text(r.get("sub_scores")),
            r.get("bonus", ""),
            dual.get("result", ""),
        ])
    for column in ws.columns:
        letter = column[0].column_letter
        width = min(max(len(str(cell.value or "")) for cell in column) + 2, 60)
        ws.column_dimensions[letter].width = width
    wb.save(target)
    return target


def export_pdf(path: Path | None = None) -> Path:
    # PDF 中文字体依赖较重；当前先导出为浏览器可打印成 PDF 的 HTML，避免乱码。
    target = path or DATA_DIR / f"评阅历史_{time.strftime('%Y%m%d_%H%M%S')}.html"
    return export_html(target)


def _esc(value: Any) -> str:
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _sub_score_text(items: Any) -> str:
    # 导出时把小题表压成一行，方便 Excel 和 HTML 同时阅读。
    if not items:
        return ""
    return "; ".join(f"{x.get('label')}={x.get('score')}/{x.get('maxScore')}" for x in items)
